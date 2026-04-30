#!/usr/bin/env python3
"""Generate pipeline_findings.docx using python-docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUT = os.path.join(os.path.dirname(__file__), "pipeline_findings.docx")

doc = Document()

# ---------------------------------------------------------------------------
# Page setup — A4, 2 cm margins
# ---------------------------------------------------------------------------
for section in doc.sections:
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.5)

# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    # Ensure font is consistent
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    return p

def body(doc, text="", space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        run = p.add_run(text)
        set_run_font(run)
    return p

def rich_para(doc, parts, space_after=6):
    """parts = list of (text, bold, italic, mono)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        font = "Courier New" if mono else "Calibri"
        set_run_font(run, name=font, size=10 if mono else 11, bold=bold, italic=italic)
    return p

def bullet(doc, parts, level=0):
    """Bullet list item. parts = list of (text, bold, italic, mono)"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    if level > 0:
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(720 + level * 360))
        ind.set(qn("w:hanging"), "360")
        pPr.append(ind)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        font = "Courier New" if mono else "Calibri"
        set_run_font(run, name=font, size=10 if mono else 11, bold=bold, italic=italic)
    return p

def numbered(doc, parts):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    for text, bold, italic, mono in parts:
        run = p.add_run(text)
        font = "Courier New" if mono else "Calibri"
        set_run_font(run, name=font, size=10 if mono else 11, bold=bold, italic=italic)
    return p

def spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)

def add_table(doc, headers, rows, col_widths_cm):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    hdr.height = Cm(0.7)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Cm(col_widths_cm[i])
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "2E74B5")
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, bold=True, color=(255, 255, 255))
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = Cm(col_widths_cm[ci])
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, size=10)
    return table

# ---------------------------------------------------------------------------
# Title block
# ---------------------------------------------------------------------------
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_run = title_p.add_run("Laser Controller — Segmentation & Vectorization Pipeline")
set_run_font(title_run, size=20, bold=True, color=(0x1F, 0x49, 0x7D))
title_p.paragraph_format.space_after = Pt(2)

sub_p = doc.add_paragraph()
sub_run = sub_p.add_run("Findings and Analysis")
set_run_font(sub_run, size=13, italic=True, color=(0x44, 0x44, 0x44))
sub_p.paragraph_format.space_after = Pt(12)

# Horizontal rule via bottom border on the subtitle paragraph
pPr = sub_p._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "2E74B5")
pBdr.append(bottom)
pPr.append(pBdr)

spacer(doc, 8)

# ---------------------------------------------------------------------------
# 1. Overview
# ---------------------------------------------------------------------------
heading(doc, "1. Overview", 1)
body(doc, ("The laser controller pipeline converts video footage into ILDA animation files "
           "suitable for real-time laser playback. The process has three stages:"))

numbered(doc, [("Segmentation", True, False, False),
               (" — isolate the subject from the video background using a neural mask.", False, False, False)])
numbered(doc, [("Vectorization", True, False, False),
               (" — convert the mask and video frame into polylines (laser beam paths).", False, False, False)])
numbered(doc, [("Encoding", True, False, False),
               (" — apply temporal filtering and write the ILDA file.", False, False, False)])

spacer(doc, 4)
body(doc, ("This document records the reasoning behind model and method choices, the shortcomings "
           "encountered, and possible directions for improvement."))
spacer(doc)

# ---------------------------------------------------------------------------
# 2. Segmentation
# ---------------------------------------------------------------------------
heading(doc, "2. Segmentation", 1)
heading(doc, "2.1  The Problem", 2)
body(doc, ("Raw video footage contains background information irrelevant to the laser show. "
           "The segmentation step produces a per-frame binary mask — pixels belonging to the "
           "subject are white, everything else is black. Downstream vectorization then operates "
           "only inside that mask region, preventing background edges from polluting the laser output."))
spacer(doc)

heading(doc, "2.2  Models Considered", 2)
body(doc, "Four segmentation model families were evaluated:")
spacer(doc, 2)

bullet(doc, [("ViLLa", True, False, False),
             (" — video instance segmentation. Strong at multi-object scenes but requires "
              "category labels and has high compute cost. Considered too heavyweight for our "
              "batch pipeline.", False, False, False)])
bullet(doc, [("DEVA", True, False, False),
             (" — decoupled video segmentation. Interesting tracking properties but less mature "
              "tooling and no simple bounding-box prompt interface.", False, False, False)])
bullet(doc, [("OMG-Seg", True, False, False),
             (" — unified segmentation model. Promising generality but complex setup and not "
              "yet widely deployed.", False, False, False)])
bullet(doc, [("Grounding DINO + SAM2", True, False, False),
             (" — a two-model stack: Grounding DINO localises a subject from a free-text prompt "
              "(e.g. “woman on motorcycle”) and returns a bounding box; SAM2 receives "
              "that box as its initial prompt and propagates a pixel-accurate mask forward "
              "through the video.", False, False, False)])
spacer(doc)

heading(doc, "2.3  Why Grounding DINO + SAM2", 2)
body(doc, "The DINO + SAM2 combination was chosen for three reasons:")
spacer(doc, 2)

numbered(doc, [("Text-guided initialisation. ", True, False, False),
               ("Grounding DINO eliminates the need to manually click a prompt point. A single "
                "descriptive phrase is enough to locate the subject in frame 0.", False, False, False)])
numbered(doc, [("Temporal propagation. ", True, False, False),
               ("SAM2 is purpose-built for video: it maintains a memory bank across frames and "
                "propagates the mask through motion, deformation, and partial occlusion without "
                "per-frame re-inference.", False, False, False)])
numbered(doc, [("Practical maturity. ", True, False, False),
               ("Both models have published checkpoints, Python packages, and active "
                "maintenance. The integration path was clear.", False, False, False)])
spacer(doc, 4)
body(doc, ("Grounding DINO runs on CPU — it processes only frame 0 and speed is not a concern. "
           "SAM2 propagation runs on CUDA. Both models are installed in a dedicated Python 3.12 "
           "virtual environment because SAM2 requires a specific PyTorch CUDA build unavailable "
           "for Python 3.14."))
spacer(doc)

heading(doc, "2.4  Limitations Observed", 2)
bullet(doc, [("Whole-scene capture. ", True, False, False),
             ("When two subjects are in close physical contact — e.g. a face with hands pressed "
              "against it — all text prompts return boxes encompassing both subjects. SAM2 then "
              "masks both as a single region. This is a fundamental limitation: SAM2 is an "
              "object tracker, not a body-part segmenter.", False, False, False)])
bullet(doc, [("Distant or low-contrast subjects. ", True, False, False),
             ("Testing on a gymnastics video with a small, dark-clothed gymnast produced a "
              "solid outline mask, but interior detail was sparse. The segmenter excels when "
              "the subject fills the frame with clear contrast.", False, False, False)])
bullet(doc, [("Source material dependency. ", True, False, False),
             ("Mask quality is highly sensitive to the source video. Subjects that fill the "
              "frame, have clear contrast against the background, and move without strong "
              "occlusion produce the best results. Source material selection is therefore "
              "considered part of the production workflow.", False, False, False)])
spacer(doc)

# ---------------------------------------------------------------------------
# 3. Vectorization
# ---------------------------------------------------------------------------
heading(doc, "3. Vectorization", 1)
heading(doc, "3.1  The Problem", 2)
body(doc, ("A binary mask defines where the subject is, not what to draw. The vectorization "
           "step converts per-frame pixel data into ordered lists of 2D points (polylines) "
           "that the laser can trace. Each polyline becomes a beam-on path; gaps between "
           "polylines are blank (beam-off) jumps. The quality of the laser image is determined "
           "almost entirely by vectorization quality."))
spacer(doc)

heading(doc, "3.2  Methods Implemented", 2)
body(doc, "Four interior edge methods are implemented, each producing a different character of line:")
spacer(doc, 2)

for label, desc in [
    ("Canny",
     "Classical gradient-based edge detection. Fast, no model required. In testing it produced "
     "extremely noisy output — high point counts, many short disconnected fragments, and heavy "
     "sensitivity to texture and compression artefacts. Not suitable as a primary method."),
    ("HED (Holistically-Nested Edge Detection)",
     "A convolutional neural network trained to predict perceptually meaningful edges. Runs via "
     "ONNX Runtime on the project’s existing HED checkpoint. In testing this produced the "
     "most visually coherent and stable lines — edges aligned with genuine contours of the "
     "subject rather than pixel-level noise. HED is the recommended primary method."),
    ("Depth",
     "A monocular depth estimation model (MiDaS/DPT) run via ONNX. Sobel gradients on the depth "
     "map locate depth discontinuities — roughly, silhouette and fore/background boundaries. "
     "Stable between frames but less visually informative than HED for interior detail. Useful as "
     "a complement or fallback."),
    ("Thin (Skeletonisation)",
     "Applies edge magnitude thresholding then Zhang-Suen thinning to produce 1-pixel-wide "
     "centerlines. In testing this produced unusual, somewhat unstable line patterns. Can capture "
     "interior structure that contour-based methods miss, but requires careful threshold tuning."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(label)
    set_run_font(r, bold=True)
    p2 = body(doc, desc, space_after=8)

spacer(doc)

heading(doc, "3.3  HED Post-Processing: Skeletonisation", 2)
rich_para(doc, [
    ("HED produces ", False, False, False),
    ("thick", False, True, False),
    (" edge blobs, not single-pixel lines. Running ", False, False, False),
    ("findContours", False, False, True),
    (" directly on the thresholded HED output traces the ", False, False, False),
    ("outline", False, True, False),
    (" of those blobs, giving double-line artefacts and closed loops around what should be "
     "single strokes.", False, False, False),
])
spacer(doc, 3)
body(doc, ("The fix applied in the current implementation is to skeletonise the HED binary mask "
           "before tracing. This collapses thick blobs to 1-pixel centerlines; the "
           "skeleton_to_polylines function then extracts ordered strokes. The result is clean "
           "single-line output that closely follows the perceptual edges HED detected."))
spacer(doc)

heading(doc, "3.4  Outer Contour", 2)
body(doc, ("Regardless of which interior method is used, the SAM2 mask boundary is always "
           "extracted as a closed outer contour. This gives the subject silhouette — the "
           "most important structural line for a laser show. The outer contour is simplified "
           "with Douglas-Peucker and drawn in every frame."))
spacer(doc)

# ---------------------------------------------------------------------------
# 4. Point optimisation
# ---------------------------------------------------------------------------
heading(doc, "4. Point Optimisation — Where, Why, and Can It Be Improved?", 1)
heading(doc, "4.1  Where It Happens", 2)
body(doc, "Point optimisation currently occurs in two stages:")
spacer(doc, 2)

bullet(doc, [("Vectorize stage. ", True, False, False),
             ("Douglas-Peucker simplification (", False, False, False),
             ("cv2.approxPolyDP", False, False, True),
             (") is applied to every polyline during extraction, controlled by per-method "
              "––epsilon parameters (default 3.0 px). A point budget (default 800 "
              "points/frame) is then enforced: if the total point count exceeds the budget, "
              "epsilon is iteratively doubled until the frame fits.", False, False, False)])
bullet(doc, [("Encode stage. ", True, False, False),
             ("The temporal persistence filter removes polylines that do not appear in a "
              "consistent spatial location across consecutive frames, reducing the total "
              "polyline count by eliminating transient noise.", False, False, False)])
spacer(doc)

heading(doc, "4.2  Why It Is Done This Way", 2)
body(doc, "The split reflects two fundamentally different kinds of problem:")
spacer(doc, 2)
bullet(doc, [("Geometric simplification", True, False, False),
             (" (Douglas-Peucker) belongs in the vectorize stage because it operates on "
              "individual polyline shapes. Doing it there avoids storing redundant points "
              "in the intermediate JSON.", False, False, False)])
bullet(doc, [("Temporal filtering", True, False, False),
             (" belongs in the encode stage because it requires comparing polylines across "
              "frames — information only available once all frames are vectorized.", False, False, False)])
spacer(doc, 4)
body(doc, ("The point budget is enforced at vectorize time rather than encode time because it "
           "is a geometric constraint (how many points the laser can draw in one frame period), "
           "not a temporal one."))
spacer(doc)

heading(doc, "4.3  Can It Be Improved?", 2)
body(doc, "Yes, significantly. Current limitations:")
spacer(doc, 2)

bullet(doc, [("Uniform budget enforcement. ", True, False, False),
             ("The current approach doubles epsilon globally until the budget is met, "
              "discarding points equally from all polylines including the outer contour, which "
              "should be protected. A better approach: simplify small interior strokes first "
              "and preserve the outer contour until last.", False, False, False)])
bullet(doc, [("No perceptual weighting. ", True, False, False),
             ("All polylines are treated equally. A long stroke describing a major edge is "
              "more visually important than ten short interior strokes. Ranking by length or "
              "salience before budget enforcement would preserve quality better.", False, False, False)])
bullet(doc, [("Static epsilon. ", True, False, False),
             ("The epsilon value is fixed per run. Adaptive epsilon — tied to local "
              "stroke curvature — would simplify straight sections aggressively while "
              "preserving curved regions.", False, False, False)])
bullet(doc, [("No point redistribution. ", True, False, False),
             ("After simplification, points are not redistributed along the stroke. Dense "
              "clustering at corners causes visible bright spots (dwell intensity). A "
              "resampling pass after simplification would address this and produce more "
              "uniform beam brightness.", False, False, False)])
spacer(doc)

# ---------------------------------------------------------------------------
# 5. Visual noise
# ---------------------------------------------------------------------------
heading(doc, "5. Visual Noise and Line Instability", 1)
heading(doc, "5.1  The Problem", 2)
body(doc, ("Even with HED producing semantically meaningful edges, the laser output can appear "
           "noisy: lines flicker between frames, short spurious strokes appear and disappear, "
           "and path geometry shifts slightly even when the subject is barely moving."))
spacer(doc)

heading(doc, "5.2  Current Mitigations", 2)
body(doc, "Three techniques are currently applied:")
spacer(doc, 2)

for label, desc in [
    ("Temporal frame blending (vectorize stage).",
     "Signal maps fed into edge detectors are exponentially blended with the previous "
     "frame’s maps (default alpha 0.7: 70% current / 30% previous). This suppresses "
     "single-frame spikes in the edge signal at the cost of a slight temporal lag on fast motion."),
    ("Minimum stroke length pruning (vectorize stage).",
     "Skeleton strokes shorter than --hed-min-len pixels (default 10 px) are discarded. "
     "This removes short dead-end branches from skeleton noise at the cost of losing genuine "
     "fine detail."),
    ("Temporal persistence filter (encode stage).",
     "A polyline is kept in frame N only if a spatially close polyline (centroid within "
     "--persist-dist pixels, default 15) existed in each of the previous --persist-frames "
     "frames (default 2). This requires a stroke to survive at least two consecutive frames "
     "before it is drawn, eliminating single-frame flashes — the most visually "
     "disruptive artefact."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(label)
    set_run_font(r, bold=True)
    body(doc, desc, space_after=8)

spacer(doc)

heading(doc, "5.3  Possible Further Improvements", 2)
spacer(doc, 2)

for label, desc in [
    ("Polyline smoothing across time.",
     "Apply a Gaussian or moving-average filter to vertex positions of each polyline across "
     "frames (smoothing the trajectory of each point, not the shape of the stroke). This "
     "requires matching corresponding polylines between frames — solvable with the "
     "centroid-distance matching already used for persistence filtering."),
    ("Morphological closing before skeletonisation.",
     "Applying dilation followed by erosion to the HED binary before skeletonising fills "
     "small gaps and merges nearby parallel edges into single strokes, reducing total stroke "
     "count and improving stability."),
    ("Edge NMS (non-maximum suppression).",
     "HED does not apply NMS by default — the output contains soft, multi-pixel-wide "
     "responses. Applying gradient direction-aware NMS (as in Canny) before thresholding "
     "produces sharper, more stable edge localisation than simple threshold + skeletonise."),
    ("Optical flow-assisted tracking.",
     "Use optical flow (e.g. Farneback or RAFT) to warp the previous frame’s polylines "
     "into the current frame as an initialisation. Strokes would move smoothly rather than "
     "being independently re-detected each frame, dramatically improving temporal consistency."),
    ("Source material selection.",
     "The single most effective quality improvement. Footage with a stationary camera, "
     "high-contrast subject, smooth background, and deliberate choreography produces clean, "
     "stable vector output. Handheld footage with compression artefacts, complex backgrounds, "
     "or fast camera motion will always challenge any vectorization pipeline regardless of "
     "algorithm choice."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(label)
    set_run_font(r, bold=True)
    body(doc, desc, space_after=8)

spacer(doc)

# ---------------------------------------------------------------------------
# 6. Summary table
# ---------------------------------------------------------------------------
heading(doc, "6. Summary", 1)
spacer(doc, 4)

add_table(doc,
    headers=["Component", "Current Choice", "Status"],
    rows=[
        ["Segmentation",    "Grounding DINO + SAM2",                    "Working; limited by source material"],
        ["Interior edges",  "HED (primary), Depth, Canny, Thin",        "HED recommended; others available"],
        ["Simplification",  "Douglas-Peucker + point budget",           "Working; improvements possible"],
        ["Noise filtering", "Frame blend + min-length + persistence",   "Working; temporal tracking not yet implemented"],
        ["ILDA output",     "Format 5, white, closed/open polylines",   "Working"],
    ],
    col_widths_cm=[3.8, 7.0, 5.6]
)

spacer(doc, 8)
body(doc, ("The pipeline is functional end-to-end. The dominant quality bottleneck is now the "
           "vectorization step, specifically the temporal stability of interior strokes. The "
           "outer contour (silhouette) is reliable; the interior detail is the open problem."))

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(OUT)
print(f"Saved: {OUT}")
